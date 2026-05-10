from docx import Document
import sys

try:
    doc = Document(r'C:\Users\Chege\Documents\forex\forex learnings\3 CANDLESTICK REVERSAL PROTOCOL.docx')
    print('Loaded successfully')
    print('Paragraphs:', len(doc.paragraphs))
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f'Para {i}: {para.text}')
except Exception as e:
    print(f'Error: {e}')
    sys.exit(1)
